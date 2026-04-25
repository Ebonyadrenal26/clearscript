"""Tests for the DOCX ingest adapter.

We construct .docx files in-memory using python-docx and feed them through
the adapter — this verifies the bold-speaker heuristic and timestamp
stripping work without needing real Feishu Miaoji exports as fixtures.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document

from clearscript.ingest import parse


def make_simple_docx(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("[00:00:00] Speaker 1: Hi everyone, can you hear me?")
    doc.add_paragraph("Speaker 2: Yes I can.")
    doc.add_paragraph("Speaker 1: Great, let's begin.")
    p = tmp_path / "simple.docx"
    doc.save(p)
    return p


def make_feishu_style_docx(tmp_path: Path) -> Path:
    """Mimic Feishu Miaoji: speaker name as a bold leading run, content unbold."""
    doc = Document()

    para1 = doc.add_paragraph()
    speaker_run = para1.add_run("张三")
    speaker_run.bold = True
    para1.add_run("：你好，能听到吗？")

    para2 = doc.add_paragraph()
    s2 = para2.add_run("李四")
    s2.bold = True
    para2.add_run("：能听到。")

    p = tmp_path / "feishu.docx"
    doc.save(p)
    return p


def test_parses_simple_docx_with_inline_speaker(tmp_path: Path) -> None:
    path = make_simple_docx(tmp_path)
    transcript = parse(path)
    assert transcript.source_format == "docx"
    # 3 turns: Speaker 1 → Speaker 2 → Speaker 1
    assert len(transcript.segments) == 3
    assert [s.speaker_raw for s in transcript.segments] == [
        "Speaker 1",
        "Speaker 2",
        "Speaker 1",
    ]


def test_parses_feishu_bold_speaker_format(tmp_path: Path) -> None:
    path = make_feishu_style_docx(tmp_path)
    transcript = parse(path)
    assert transcript.source_format == "docx"
    speakers = [s.speaker_raw for s in transcript.segments]
    assert "张三" in speakers
    assert "李四" in speakers


def test_strips_leading_timestamp_in_paragraph(tmp_path: Path) -> None:
    doc = Document()
    doc.add_paragraph("[0:14:33] Speaker 1: Hello there.")
    p = tmp_path / "ts.docx"
    doc.save(p)

    transcript = parse(p)
    assert transcript.segments[0].speaker_raw == "Speaker 1"
    assert "Hello there" in transcript.segments[0].text
    assert "0:14:33" not in transcript.segments[0].text


def test_empty_docx_returns_empty_transcript(tmp_path: Path) -> None:
    doc = Document()
    p = tmp_path / "empty.docx"
    doc.save(p)
    transcript = parse(p)
    assert transcript.segments == []
